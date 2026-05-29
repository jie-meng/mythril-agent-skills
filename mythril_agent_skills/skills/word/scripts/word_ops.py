#!/usr/bin/env python3
"""Word document operations CLI for AI coding assistants.

Provides subcommands for reading, extracting, and inspecting Word documents.

Requires: python-docx (pip install python-docx)
Optional: python-docx for .docx, antiword CLI for .doc
"""

from __future__ import annotations

import argparse
import csv
import json
import sys
from pathlib import Path
from typing import Any

try:
    from docx import Document
    from docx.opc.constants import RELATIONSHIP_TYPE as RT
except ImportError:
    print(
        "ERROR: python-docx is not installed.\n"
        "Install it with: pip install python-docx",
        file=sys.stderr,
    )
    sys.exit(1)


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def _format_size(size_bytes: int) -> str:
    """Format file size in human-readable form."""
    if size_bytes < 1024:
        return f"{size_bytes} B"
    if size_bytes < 1024 * 1024:
        return f"{size_bytes / 1024:.1f} KB"
    return f"{size_bytes / (1024 * 1024):.1f} MB"


def _load_docx(path: str) -> Document:
    """Load a .docx document, handling common errors."""
    p = Path(path)
    if not p.exists():
        print(f"ERROR: File not found: {path}", file=sys.stderr)
        sys.exit(2)
    if p.suffix.lower() not in (".docx", ".docm"):
        print(
            f"ERROR: Not a Word document (.docx/.docm): {path}\n"
            "For legacy .doc files, convert to .docx first:\n"
            "  - LibreOffice: libreoffice --headless --convert-to docx file.doc\n"
            "  - Or use an online converter",
            file=sys.stderr,
        )
        sys.exit(2)
    try:
        return Document(str(p))
    except Exception as exc:
        print(f"ERROR: Cannot read Word document: {exc}", file=sys.stderr)
        sys.exit(2)
        raise  # unreachable, satisfies type checker


def _extract_paragraph_text(para: Any) -> str:
    """Extract text from a paragraph, including runs."""
    return para.text


def _table_to_markdown(table: Any) -> str:
    """Convert a docx Table to markdown format."""
    rows: list[list[str]] = []
    for row in table.rows:
        cells = [cell.text.strip().replace("|", "\\|") for cell in row.cells]
        rows.append(cells)

    if not rows:
        return ""

    # Determine column count from max row
    max_cols = max(len(r) for r in rows)

    # Pad rows to uniform width
    for row in rows:
        while len(row) < max_cols:
            row.append("")

    lines: list[str] = []
    # Header
    lines.append("| " + " | ".join(rows[0]) + " |")
    lines.append("| " + " | ".join("---" for _ in rows[0]) + " |")
    # Body
    for row in rows[1:]:
        lines.append("| " + " | ".join(row) + " |")

    return "\n".join(lines)


def _table_to_json(table: Any) -> list[dict[str, str]]:
    """Convert a docx Table to a list of dicts (first row = keys)."""
    rows: list[list[str]] = []
    for row in table.rows:
        rows.append([cell.text.strip() for cell in row.cells])

    if len(rows) < 2:
        return []

    headers = rows[0]
    result: list[dict[str, str]] = []
    for row in rows[1:]:
        record: dict[str, str] = {}
        for i, header in enumerate(headers):
            record[header] = row[i] if i < len(row) else ""
        result.append(record)
    return result


def _count_words(text: str) -> int:
    """Count words in text (handles CJK characters as individual words)."""
    import re

    # Split CJK characters as individual tokens, then count
    cjk = len(re.findall(r"[\u4e00-\u9fff\u3400-\u4dbf\uf900-\ufaff]", text))
    # Count non-CJK words
    non_cjk_text = re.sub(r"[\u4e00-\u9fff\u3400-\u4dbf\uf900-\ufaff]", " ", text)
    non_cjk_words = len(non_cjk_text.split())
    return cjk + non_cjk_words


# ---------------------------------------------------------------------------
# Subcommands
# ---------------------------------------------------------------------------


def cmd_info(args: argparse.Namespace) -> None:
    """Show Word document metadata and structure."""
    doc = _load_docx(args.file)
    p = Path(args.file)

    print(f"# Word Document Info: {p.name}\n")
    print(f"- **File size**: {_format_size(p.stat().st_size)}")

    # Core properties
    props = doc.core_properties
    if props.title:
        print(f"- **Title**: {props.title}")
    if props.author:
        print(f"- **Author**: {props.author}")
    if props.subject:
        print(f"- **Subject**: {props.subject}")
    if props.keywords:
        print(f"- **Keywords**: {props.keywords}")
    if props.comments:
        print(f"- **Comments**: {props.comments}")
    if props.created:
        print(f"- **Created**: {props.created}")
    if props.modified:
        print(f"- **Modified**: {props.modified}")
    if props.last_modified_by:
        print(f"- **Last modified by**: {props.last_modified_by}")
    if props.revision:
        print(f"- **Revision**: {props.revision}")
    if props.category:
        print(f"- **Category**: {props.category}")
    if props.content_status:
        print(f"- **Status**: {props.content_status}")

    # Document structure
    para_count = len(doc.paragraphs)
    table_count = len(doc.tables)

    # Count words across all paragraphs
    all_text = "\n".join(p.text for p in doc.paragraphs)
    word_count = _count_words(all_text)

    print(f"\n## Structure\n")
    print(f"- **Paragraphs**: {para_count}")
    print(f"- **Tables**: {table_count}")
    print(f"- **Words**: {word_count}")

    # Section info
    section_count = len(doc.sections)
    print(f"- **Sections**: {section_count}")

    # Heading outline
    headings: list[tuple[int, str]] = []
    for para in doc.paragraphs:
        if para.style and para.style.name and para.style.name.startswith("Heading"):
            try:
                level = int(para.style.name.replace("Heading ", ""))
            except ValueError:
                continue
            text = para.text.strip()
            if text:
                headings.append((level, text))

    if headings:
        print(f"\n## Headings\n")
        for level, text in headings[:50]:  # Limit to first 50
            indent = "  " * (level - 1)
            print(f"{indent}- {text}")


def cmd_text(args: argparse.Namespace) -> None:
    """Extract text from Word document."""
    doc = _load_docx(args.file)

    paragraphs = doc.paragraphs
    if args.limit:
        paragraphs = paragraphs[: args.limit]

    for i, para in enumerate(paragraphs):
        text = para.text
        if text.strip():
            # Show heading style indicator
            style_name = para.style.name if para.style else ""
            if style_name.startswith("Heading"):
                level = style_name.replace("Heading ", "")
                print(f"{'#' * int(level)} {text}")
                print()
            else:
                print(text)
                print()
        elif not args.skip_empty:
            print()


def cmd_tables(args: argparse.Namespace) -> None:
    """Extract tables from Word document."""
    doc = _load_docx(args.file)

    if not doc.tables:
        print("No tables found in the document.")
        return

    fmt = args.format

    if fmt == "json":
        all_tables: list[dict[str, Any]] = []
        for i, table in enumerate(doc.tables):
            rows = _table_to_json(table)
            if rows:
                all_tables.append({"table_number": i + 1, "rows": rows})
        print(json.dumps(all_tables, indent=2, ensure_ascii=False))

    elif fmt == "csv":
        output_dir = Path(args.output_dir) if args.output_dir else Path(".")
        output_dir.mkdir(parents=True, exist_ok=True)
        stem = Path(args.file).stem

        for i, table in enumerate(doc.tables):
            csv_path = output_dir / f"{stem}_table_{i + 1}.csv"
            with open(csv_path, "w", newline="", encoding="utf-8") as f:
                writer = csv.writer(f)
                for row in table.rows:
                    writer.writerow([cell.text.strip() for cell in row.cells])
            print(f"Saved: {csv_path}")

    else:  # markdown
        for i, table in enumerate(doc.tables):
            md = _table_to_markdown(table)
            if md:
                print(f"## Table {i + 1}\n")
                print(md)
                print()


def cmd_extract_images(args: argparse.Namespace) -> None:
    """Extract embedded images from Word document."""
    doc = _load_docx(args.file)
    output_dir = Path(args.output_dir) if args.output_dir else Path(".")
    output_dir.mkdir(parents=True, exist_ok=True)
    stem = Path(args.file).stem

    img_count = 0
    for rel in doc.part.rels.values():
        if "image" in rel.reltype:
            img_count += 1
            ext = rel.target_ref.split(".")[-1] if "." in rel.target_ref else "bin"
            out_path = output_dir / f"{stem}_image_{img_count}.{ext}"
            with open(out_path, "wb") as f:
                f.write(rel.target_part.blob)
            print(f"Extracted: {out_path}")

    if img_count == 0:
        print("No embedded images found in the document.")
    else:
        print(f"\nExtracted {img_count} image(s).")


def cmd_to_markdown(args: argparse.Namespace) -> None:
    """Convert Word document to Markdown format."""
    doc = _load_docx(args.file)

    lines: list[str] = []
    table_index = 0

    for element in doc.element.body:
        tag = element.tag.split("}")[-1] if "}" in element.tag else element.tag

        if tag == "p":
            # Paragraph
            para = None
            for p in doc.paragraphs:
                if p._element is element:
                    para = p
                    break
            if para is None:
                continue

            text = para.text.strip()
            if not text:
                lines.append("")
                continue

            style_name = para.style.name if para.style else ""
            if style_name.startswith("Heading"):
                try:
                    level = int(style_name.replace("Heading ", ""))
                except ValueError:
                    level = 1
                lines.append(f"{'#' * level} {text}")
                lines.append("")
            else:
                # Check for bold/italic in runs
                formatted_parts: list[str] = []
                for run in para.runs:
                    run_text = run.text
                    if not run_text:
                        continue
                    if run.bold and run.italic:
                        formatted_parts.append(f"***{run_text}***")
                    elif run.bold:
                        formatted_parts.append(f"**{run_text}**")
                    elif run.italic:
                        formatted_parts.append(f"*{run_text}*")
                    else:
                        formatted_parts.append(run_text)
                lines.append("".join(formatted_parts))
                lines.append("")

        elif tag == "tbl":
            # Table
            if table_index < len(doc.tables):
                table = doc.tables[table_index]
                table_index += 1
                md = _table_to_markdown(table)
                if md:
                    lines.append(md)
                    lines.append("")

    output = "\n".join(lines)

    if args.output:
        out_path = Path(args.output)
        out_path.write_text(output, encoding="utf-8")
        print(f"Saved: {out_path}")
    else:
        print(output)


# ---------------------------------------------------------------------------
# CLI parser
# ---------------------------------------------------------------------------


def build_parser() -> argparse.ArgumentParser:
    """Build the argument parser with all subcommands."""
    parser = argparse.ArgumentParser(
        prog="word_ops.py",
        description="Word document operations for AI coding assistants.",
    )
    sub = parser.add_subparsers(dest="command", required=True)

    # info
    p_info = sub.add_parser("info", help="Show document metadata and structure")
    p_info.add_argument("file", help="Word document path (.docx/.docm)")

    # text
    p_text = sub.add_parser("text", help="Extract text from document")
    p_text.add_argument("file", help="Word document path (.docx/.docm)")
    p_text.add_argument(
        "--limit", type=int, help="Limit number of paragraphs"
    )
    p_text.add_argument(
        "--skip-empty",
        action="store_true",
        help="Skip empty paragraphs",
    )

    # tables
    p_tables = sub.add_parser("tables", help="Extract tables from document")
    p_tables.add_argument("file", help="Word document path (.docx/.docm)")
    p_tables.add_argument(
        "--format",
        choices=["markdown", "json", "csv"],
        default="markdown",
        help="Output format (default: markdown)",
    )
    p_tables.add_argument(
        "--output-dir", help="Output directory for CSV files"
    )

    # extract-images
    p_eimg = sub.add_parser(
        "extract-images", help="Extract embedded images from document"
    )
    p_eimg.add_argument("file", help="Word document path (.docx/.docm)")
    p_eimg.add_argument("--output-dir", help="Output directory for images")

    # to-markdown
    p_md = sub.add_parser("to-markdown", help="Convert document to Markdown")
    p_md.add_argument("file", help="Word document path (.docx/.docm)")
    p_md.add_argument("-o", "--output", help="Output file path")

    return parser


COMMAND_DISPATCH: dict[str, Any] = {
    "info": cmd_info,
    "text": cmd_text,
    "tables": cmd_tables,
    "extract-images": cmd_extract_images,
    "to-markdown": cmd_to_markdown,
}


def main() -> None:
    """Entry point."""
    parser = build_parser()
    args = parser.parse_args()
    handler = COMMAND_DISPATCH.get(args.command)
    if handler:
        handler(args)
    else:
        parser.print_help()
        sys.exit(1)


if __name__ == "__main__":
    main()
